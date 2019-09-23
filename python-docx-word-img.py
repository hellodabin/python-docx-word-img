# coding:utf-8
# 导入Python-docx相关的包
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
# 解析路径用
import os

# 生成文档对象
document = Document()
# 因为现在不支持直接图片定位，所以需要用空文本替换的方式来解决图片方向问题，生成一个段落
paragraph = document.add_paragraph()

# 图片向左设置
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = paragraph.add_run("")

# 图片文件夹路径，需要的自行修改
path = "/Users/dabin/Downloads/img"
img_paths = os.listdir(path)
# 批量生成图片路径
for img_path in img_paths:
    if img_path != ".DS_Store":
        # 空文字插入图片替换，达到图片居左不换行的效果
        run.add_picture("{}/{}".format(path, img_path),
                        width=Inches(1.5))
# 保存文档
document.save("img.docx")
